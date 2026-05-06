from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
PROMO_DIR = ROOT / "data" / "raw" / "promociones"
OUTPUT_PATH = ROOT / "data" / "interim" / "promociones_extraidas.xlsx"


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            table_cells.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(paragraphs + table_cells).strip()


def list_promo_files() -> list[Path]:
    extensions = {".pdf", ".docx"}
    return sorted(path for path in PROMO_DIR.rglob("*") if path.suffix.lower() in extensions)


def main() -> None:
    files = list_promo_files()
    if not files:
        print("No encontre PDF o DOCX en data/raw/promociones.")
        return

    rows = []
    for path in files:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = extract_pdf_text(path)
        elif suffix == ".docx":
            text = extract_docx_text(path)
        else:
            continue

        rows.append(
            {
                "archivo": str(path.relative_to(ROOT)),
                "tipo": suffix.replace(".", ""),
                "texto_extraido": text,
                "fecha_inicio": "",
                "fecha_fin": "",
                "producto": "",
                "descuento": "",
                "mecanica": "",
                "notas_revision": "",
            }
        )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(rows).to_excel(OUTPUT_PATH, index=False)
    print(f"Promociones extraidas: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

